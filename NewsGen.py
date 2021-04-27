import PySimpleGUI as sg
import docx
import datetime
import time
from docx2pdf import convert
import os
import win32com.client as client
from pathlib import Path
import openpyxl
from openpyxl import load_workbook
import comtypes.client

wdFormatPDF = 17

# Creating Path
CurrentDirectory = Path.cwd()

# Creating date object
todaysdate = datetime.datetime.now()

# Accessing the excel file
wb = load_workbook('NEWS_Status.xlsx')
ws = wb['Sheet1']

# Emails HTMLs
CemaEmail = open('Toscano Files/CemaEmail.html', 'r').read()
CoopEmail = open('Toscano Files/CoopEmail.html', 'r').read()

# Accessing Outlook Application
outlook = client.Dispatch("Outlook.Application")

# Opening up the Original Invoice Sheet to modify
doc = docx.Document('Toscano Files/InvoiceSheet.docx')
CoopNotice = docx.Document('Toscano Files/Coop30DayNotice.docx')
CemaNotice = docx.Document('Toscano Files/CEMA30DayNotice.docx')
CoopScheduling = docx.Document('Toscano Files/CoopSchedulingForm.docx')
CemaScheduling = docx.Document('Toscano Files/CEMASchedulingForm.docx')
CoopUpdateSheet = docx.Document('Toscano Files/CoopUpdateSheet.docx')
CemaUpdateSheet = docx.Document('Toscano Files/CEMAUpdateSheet.docx')
CemaAssignmentForm = docx.Document('Toscano Files/AssignmentForm.docx')
CoopHELOCNotice = docx.Document('Toscano Files/HELOCCoop30DayNotice.docx')
Coop1stHELOCNotice = docx.Document('Toscano Files/1st+HELOCCoop30DayNotice.docx')
CemaHELOCNotice = docx.Document('Toscano Files/HELOCCEMA30DayNotice.docx')

# Theme COLORS!
sg.theme('GreenMono')

# All the stuff inside your window.
layout = [  [sg.Text('New File Form Generator', font = ('Calibri', 22))],
            [sg.Text('Borrower'), sg.InputText(key = '-borrower-')],
            [sg.Text('Loan Number'), sg.InputText(key = '-loannumber-')],
            [sg.Text('HELOC #'),sg.InputText(key = '-helocnumber-')],
            [sg.Text('Contact'), sg.InputText(key = '-contact-')],
            [sg.Text('Contact Email'), sg.InputText(key = '-contactemail-')],
            [sg.Text('Phone Number'), sg.InputText(key = '-phonenumber-')],
            [sg.Text('Residents'), sg.InputText(key = '-residents-')],
            [sg.Text('Street Address'), sg.InputText(key = '-streetaddress-')],
            [sg.Text('City, State Zip'), sg.InputText(key = '-citystatezip-')],
            [sg.Combo(['CEMA', 'CEMA HELOC', 'Coop', 'Coop 1st + HELOC', 'Coop HELOC'], key = 'LoanType')],
            [sg.Checkbox('Create PDFs', key = 'PDF'), sg.Checkbox('Send Notice', key = 'Notice')],
            [sg.Button('Enter'), sg.Button('Clear'), sg.Button('Exit')],
            [sg.Text('                                                                                                                                  ', key = 'Output')],
            ]

# Create the Window 
window = sg.Window('News Generator', layout, default_element_size = (60,1), location = (50,150), resizable = True)

# Event Loop to process "events" and get the "values" of the inputs
while True:
    event, values = window.read()
    if event == sg.WIN_CLOSED or event == 'Exit': # if user closes window or clicks cancel
        break
    if event == 'Enter': # and values['-borrower-'] == True and values['-loannumber-'] == True and values['-contactemail-'] == True and values['-residents-'] == True and values['-citystatezip-'] == True
        # Assigning the values taken from the fields to variables
        Borrower = values['-borrower-']
        LoanNumber = values['-loannumber-']
        HelocNumber = values['-helocnumber-']
        Contact = values['-contact-']
        ContactEmail = values['-contactemail-']
        PhoneNumber = values['-phonenumber-']
        Residents = values['-residents-']
        StreetAddress = values['-streetaddress-']
        CityStateZip = values['-citystatezip-']


        # CEMA 
        if values['LoanType'] == 'CEMA':
            # Appending to Excel
            newRowLocation = ws.max_row + 1
            ws.cell(column = 1, row = newRowLocation, value = Borrower)
            ws.cell(column = 2, row = newRowLocation, value = LoanNumber)
            ws.cell(column = 5, row = newRowLocation, value = Contact)
            ws.cell(column = 6, row = newRowLocation, value = ContactEmail)
            ws.cell(column = 4, row = newRowLocation, value = todaysdate.strftime("%m/%d/%Y"))
            wb.save('NEWS_Status.xlsx')
            wb.close()
            # Invoice
            doc.paragraphs[18].runs[2].text = Borrower
            doc.paragraphs[18].runs[2].font.name = 'Times New Roman'
            doc.paragraphs[18].runs[8].text = LoanNumber
            doc.paragraphs[18].runs[8].font.name = 'Times New Roman'
            doc.paragraphs[12].text = todaysdate.strftime("%B %d, %Y")
            doc.paragraphs[12].runs[0].font.name = 'Times New Roman'
            doc.paragraphs[13].text = Residents
            doc.paragraphs[13].runs[0].font.name = 'Times New Roman'
            doc.paragraphs[14].text = StreetAddress
            doc.paragraphs[14].runs[0].font.name = 'Times New Roman'
            doc.paragraphs[15].text = CityStateZip
            doc.paragraphs[15].runs[0].font.name = 'Times New Roman'
            doc.paragraphs[27].runs[3].text = '250'
            doc.paragraphs[27].runs[3].font.name = 'Times New Roman'
            doc.save('News/{} - {} - Invoice.docx'.format(Borrower, LoanNumber))
            if values['PDF'] == True:
                word = comtypes.client.CreateObject('Word.Application')
                tempdoc = word.Documents.Open(os.path.abspath('News/{} - {} - Invoice.docx'.format(Borrower, LoanNumber)))
                tempdoc.SaveAs(os.path.abspath('News/{} - {} - Invoice.pdf'.format(Borrower, LoanNumber)), FileFormat=wdFormatPDF)
                tempdoc.Close()
                word.Quit()
            # CEMA30DayNotice
            CemaNotice.paragraphs[9].runs[2].text = Contact
            CemaNotice.paragraphs[9].runs[2].font.name = 'Times New Roman'
            CemaNotice.paragraphs[11].runs[5].text = ContactEmail
            CemaNotice.paragraphs[11].runs[5].font.name = 'Times New Roman'
            CemaNotice.paragraphs[15].runs[2].text = todaysdate.strftime("%m/%d/%Y")
            CemaNotice.paragraphs[15].runs[2].font.name = 'Times New Roman'
            CemaNotice.paragraphs[17].runs[3].text = Borrower
            CemaNotice.paragraphs[17].runs[3].font.name = 'Times New Roman'
            CemaNotice.paragraphs[17].runs[8].text = LoanNumber
            CemaNotice.paragraphs[17].runs[8].font.name = 'Times New Roman'
            CemaNotice.save('News/{} - {} - CEMA30DayNotice.docx'.format(Borrower, LoanNumber))
            if values['PDF'] == True:
                word = comtypes.client.CreateObject('Word.Application')
                tempdoc = word.Documents.Open(os.path.abspath('News/{} - {} - CEMA30DayNotice.docx'.format(Borrower, LoanNumber)))
                tempdoc.SaveAs(os.path.abspath('News/{} - {} - CEMA30DayNotice.pdf'.format(Borrower, LoanNumber)), FileFormat=wdFormatPDF)
                tempdoc.Close()
                word.Quit()
            # CEMA Scheduling Form
            CemaScheduling.paragraphs[9].runs[2].text = Contact
            CemaScheduling.paragraphs[9].runs[2].font.name = 'Times New Roman'
            CemaScheduling.paragraphs[10].runs[7].text = ContactEmail
            CemaScheduling.paragraphs[10].runs[7].font.name = 'Times New Roman'
            CemaScheduling.paragraphs[13].runs[5].text = Borrower
            CemaScheduling.paragraphs[13].runs[5].font.name = 'Times New Roman'
            CemaScheduling.paragraphs[13].runs[11].text = LoanNumber
            CemaScheduling.paragraphs[13].runs[11].font.name = 'Times New Roman'
            CemaScheduling.save('Scheduling Forms/{} - {} - Scheduling Form.docx'.format(Borrower, LoanNumber))
            # Assignment Form
            CemaAssignmentForm.paragraphs[9].runs[2].text = Contact
            CemaAssignmentForm.paragraphs[9].runs[2].font.name = 'Times New Roman'
            CemaAssignmentForm.paragraphs[10].runs[7].text = ContactEmail
            CemaAssignmentForm.paragraphs[10].runs[7].font.name = 'Times New Roman'
            CemaAssignmentForm.paragraphs[14].runs[4].text = Borrower
            CemaAssignmentForm.paragraphs[14].runs[4].font.name = 'Times New Roman'
            CemaAssignmentForm.paragraphs[14].runs[9].text = str(LoanNumber)
            CemaAssignmentForm.paragraphs[14].runs[9].font.name = 'Times New Roman'
            CemaAssignmentForm.save('Assignment Forms/{} - {} - Assignment Form.docx'.format(Borrower, LoanNumber))
            # Update Sheet
            CemaUpdateSheet.paragraphs[0].runs[0].text = Borrower
            CemaUpdateSheet.paragraphs[0].runs[0].font.name = 'Times New Roman'
            CemaUpdateSheet.paragraphs[0].runs[5].text = LoanNumber
            CemaUpdateSheet.paragraphs[0].runs[5].font.name = 'Times New Roman'
            CemaUpdateSheet.paragraphs[16].runs[2].text = Contact
            CemaUpdateSheet.paragraphs[16].runs[2].font.name = 'Times New Roman'
            CemaUpdateSheet.paragraphs[17].runs[1].text = ContactEmail
            CemaUpdateSheet.paragraphs[17].runs[1].font.name = 'Times New Roman'
            CemaUpdateSheet.paragraphs[18].runs[1].text = PhoneNumber
            CemaUpdateSheet.paragraphs[18].runs[1].font.name = 'Times New Roman'
            CemaUpdateSheet.save('News/{} - {} - UpdateSheet.docx'.format(Borrower, LoanNumber))
            if values['PDF'] == True:
                word = comtypes.client.CreateObject('Word.Application')
                tempdoc = word.Documents.Open(os.path.abspath('News/{} - {} - UpdateSheet.docx'.format(Borrower, LoanNumber)))
                tempdoc.SaveAs(os.path.abspath('News/{} - {} - UpdateSheet.pdf'.format(Borrower, LoanNumber)), FileFormat=wdFormatPDF)
                tempdoc.Close()
                word.Quit()
            # CEMA Email with PDF Attachments
            if values['PDF'] == True and values['Notice'] == True:
                message = outlook.CreateItem(0)
                message.To = ContactEmail
                # message.CC = 'chrisheyer0@gmail.com'
                message.Subject = '{} - {} - Initial Notice'.format(Borrower, LoanNumber)
                message.HTMLBody = CemaEmail
                message.Attachments.Add('{}/News/{} - {} - Invoice.pdf'.format(CurrentDirectory, Borrower, LoanNumber))
                message.Attachments.Add('{}/News/{} - {} - CEMA30DayNotice.pdf'.format(CurrentDirectory, Borrower, LoanNumber))
                message.Send()
                window['Output'].update('CEMA Initial Notice sent to: {} - {}'.format(ContactEmail, datetime.datetime.now().strftime("%m/%d/%Y at %I:%M:%S %p")))
            # CEMA Email with DOCX Attachments
            if values['PDF'] == False and values['Notice'] == True:
                message = outlook.CreateItem(0)
                message.To = ContactEmail
                # message.CC = 'chrisheyer0@gmail.com'
                message.Subject = '{} - {} - Initial Notice'.format(Borrower, LoanNumber)
                message.HTMLBody = CemaEmail
                message.Attachments.Add('{}/News/{} - {} - Invoice.docx'.format(CurrentDirectory, Borrower, LoanNumber))
                message.Attachments.Add('{}/News/{} - {} - CEMA30DayNotice.docx'.format(CurrentDirectory, Borrower, LoanNumber))
                message.Send()
                window['Output'].update('CEMA Initial Notice sent to: {} - {}'.format(ContactEmail, datetime.datetime.now().strftime("%m/%d/%Y at %I:%M:%S %p")))
        
                
        # CEMA HELOC
        elif values['LoanType'] == 'CEMA HELOC':
            LoanNumber = HelocNumber
            # Appending to Excel
            newRowLocation = ws.max_row + 1
            ws.cell(column = 1, row = newRowLocation, value = Borrower)
            ws.cell(column = 2, row = newRowLocation, value = LoanNumber)
            ws.cell(column = 5, row = newRowLocation, value = Contact)
            ws.cell(column = 6, row = newRowLocation, value = ContactEmail)
            ws.cell(column = 4, row = newRowLocation, value = todaysdate.strftime("%m/%d/%Y"))
            wb.save('NEWS_Status.xlsx')
            wb.close()
            # Invoice
            doc.paragraphs[18].runs[2].text = Borrower
            doc.paragraphs[18].runs[2].font.name = 'Times New Roman'
            doc.paragraphs[18].runs[8].text = LoanNumber
            doc.paragraphs[18].runs[8].font.name = 'Times New Roman'
            doc.paragraphs[12].text = todaysdate.strftime("%B %d, %Y")
            doc.paragraphs[12].runs[0].font.name = 'Times New Roman'
            doc.paragraphs[13].text = Residents
            doc.paragraphs[13].runs[0].font.name = 'Times New Roman'
            doc.paragraphs[14].text = StreetAddress
            doc.paragraphs[14].runs[0].font.name = 'Times New Roman'
            doc.paragraphs[15].text = CityStateZip
            doc.paragraphs[15].runs[0].font.name = 'Times New Roman'
            doc.paragraphs[27].runs[3].text = '250'
            doc.paragraphs[27].runs[3].font.name = 'Times New Roman'
            doc.save('News/{} - {} - Invoice.docx'.format(Borrower, LoanNumber))
            if values['PDF'] == True:
                word = comtypes.client.CreateObject('Word.Application')
                tempdoc = word.Documents.Open(os.path.abspath('News/{} - {} - Invoice.docx'.format(Borrower, LoanNumber)))
                tempdoc.SaveAs(os.path.abspath('News/{} - {} - Invoice.pdf'.format(Borrower, LoanNumber)), FileFormat=wdFormatPDF)
                tempdoc.Close()
                word.Quit()
            # CEMA30DayNotice
            CemaHELOCNotice.paragraphs[9].runs[2].text = Contact
            CemaHELOCNotice.paragraphs[9].runs[2].font.name = 'Times New Roman'
            CemaHELOCNotice.paragraphs[11].runs[6].text = ContactEmail
            CemaHELOCNotice.paragraphs[11].runs[6].font.name = 'Times New Roman'
            CemaHELOCNotice.paragraphs[15].runs[2].text = todaysdate.strftime("%m/%d/%Y")
            CemaHELOCNotice.paragraphs[15].runs[2].font.name = 'Times New Roman'
            CemaHELOCNotice.paragraphs[17].runs[3].text = Borrower
            CemaHELOCNotice.paragraphs[17].runs[3].font.name = 'Times New Roman'
            CemaHELOCNotice.paragraphs[17].runs[8].text = LoanNumber
            CemaHELOCNotice.paragraphs[17].runs[8].font.name = 'Times New Roman'
            CemaHELOCNotice.save('News/{} - {} - CEMA30DayNotice.docx'.format(Borrower, LoanNumber))
            if values['PDF'] == True:
                word = comtypes.client.CreateObject('Word.Application')
                tempdoc = word.Documents.Open(os.path.abspath('News/{} - {} - CEMA30DayNotice.docx'.format(Borrower, LoanNumber)))
                tempdoc.SaveAs(os.path.abspath('News/{} - {} - CEMA30DayNotice.pdf'.format(Borrower, LoanNumber)), FileFormat=wdFormatPDF)
                tempdoc.Close()
                word.Quit() 
            # CEMA Scheduling Form
            CemaScheduling.paragraphs[9].runs[2].text = Contact
            CemaScheduling.paragraphs[9].runs[2].font.name = 'Times New Roman'
            CemaScheduling.paragraphs[10].runs[7].text = ContactEmail
            CemaScheduling.paragraphs[10].runs[7].font.name = 'Times New Roman'
            CemaScheduling.paragraphs[13].runs[5].text = Borrower
            CemaScheduling.paragraphs[13].runs[5].font.name = 'Times New Roman'
            CemaScheduling.paragraphs[13].runs[11].text = LoanNumber
            CemaScheduling.paragraphs[13].runs[11].font.name = 'Times New Roman'
            CemaScheduling.save('Scheduling Forms/{} - {} - Scheduling Form.docx'.format(Borrower, LoanNumber))
            # Assignment Form
            CemaAssignmentForm.paragraphs[9].runs[2].text = Contact
            CemaAssignmentForm.paragraphs[9].runs[2].font.name = 'Times New Roman'
            CemaAssignmentForm.paragraphs[10].runs[7].text = ContactEmail
            CemaAssignmentForm.paragraphs[10].runs[7].font.name = 'Times New Roman'
            CemaAssignmentForm.paragraphs[14].runs[4].text = Borrower
            CemaAssignmentForm.paragraphs[14].runs[4].font.name = 'Times New Roman'
            CemaAssignmentForm.paragraphs[14].runs[9].text = str(LoanNumber)
            CemaAssignmentForm.paragraphs[14].runs[9].font.name = 'Times New Roman'
            CemaAssignmentForm.save('Assignment Forms/{} - {} - Assignment Form.docx'.format(Borrower, LoanNumber))
            # Update Sheet
            CemaUpdateSheet.paragraphs[0].runs[0].text = Borrower
            CemaUpdateSheet.paragraphs[0].runs[0].font.name = 'Times New Roman'
            CemaUpdateSheet.paragraphs[0].runs[5].text = LoanNumber
            CemaUpdateSheet.paragraphs[0].runs[5].font.name = 'Times New Roman'
            CemaUpdateSheet.paragraphs[16].runs[2].text = Contact
            CemaUpdateSheet.paragraphs[16].runs[2].font.name = 'Times New Roman'
            CemaUpdateSheet.paragraphs[17].runs[1].text = ContactEmail
            CemaUpdateSheet.paragraphs[17].runs[1].font.name = 'Times New Roman'
            CemaUpdateSheet.paragraphs[18].runs[1].text = PhoneNumber
            CemaUpdateSheet.paragraphs[18].runs[1].font.name = 'Times New Roman'
            CemaUpdateSheet.save('News/{} - {} - UpdateSheet.docx'.format(Borrower, LoanNumber))
            if values['PDF'] == True:
                word = comtypes.client.CreateObject('Word.Application')
                tempdoc = word.Documents.Open(os.path.abspath('News/{} - {} - UpdateSheet.docx'.format(Borrower, LoanNumber)))
                tempdoc.SaveAs(os.path.abspath('News/{} - {} - UpdateSheet.pdf'.format(Borrower, LoanNumber)), FileFormat=wdFormatPDF)
                tempdoc.Close()
                word.Quit()
            # CEMA Email with PDF Attachments
            if values['PDF'] == True and values['Notice'] == True:
                message = outlook.CreateItem(0)
                message.To = ContactEmail
                # message.CC = 'chrisheyer0@gmail.com'
                message.Subject = '{} - {} - Initial Notice'.format(Borrower, LoanNumber)
                message.HTMLBody = CemaEmail
                message.Attachments.Add('{}/News/{} - {} - Invoice.pdf'.format(CurrentDirectory, Borrower, LoanNumber))
                message.Attachments.Add('{}/News/{} - {} - CEMA30DayNotice.pdf'.format(CurrentDirectory, Borrower, LoanNumber))
                message.Send()
                window['Output'].update('CEMA Initial Notice sent to: {} - {}'.format(ContactEmail, datetime.datetime.now().strftime("%m/%d/%Y at %I:%M:%S %p")))
                # CEMA Email with DOCX Attachments
            if values['PDF'] == False and values['Notice'] == True:
                message = outlook.CreateItem(0)
                message.To = ContactEmail
                # message.CC = 'chrisheyer0@gmail.com'
                message.Subject = '{} - {} - Initial Notice'.format(Borrower, LoanNumber)
                message.HTMLBody = CemaEmail
                message.Attachments.Add('{}/News/{} - {} - Invoice.docx'.format(CurrentDirectory, Borrower, LoanNumber))
                message.Attachments.Add('{}/News/{} - {} - CEMA30DayNotice.docx'.format(CurrentDirectory, Borrower, LoanNumber))
                message.Send()
                window['Output'].update('CEMA HELOC Initial Notice sent to: {} - {}'.format(ContactEmail, datetime.datetime.now().strftime("%m/%d/%Y at %I:%M:%S %p")))


        # COOP
        elif values['LoanType'] == 'Coop':
            # Appending to Excel
            newRowLocation = ws.max_row + 1
            ws.cell(column = 1, row = newRowLocation, value = Borrower)
            ws.cell(column = 2, row = newRowLocation, value = LoanNumber)
            ws.cell(column = 5, row = newRowLocation, value = Contact)
            ws.cell(column = 6, row = newRowLocation, value = ContactEmail)
            ws.cell(column = 3, row = newRowLocation, value = todaysdate.strftime("%m/%d/%Y"))
            wb.save('NEWS_Status.xlsx')
            wb.close()
            # Invoice
            doc.paragraphs[18].runs[2].text = Borrower
            doc.paragraphs[18].runs[2].font.name = 'Times New Roman'
            doc.paragraphs[18].runs[8].text = LoanNumber
            doc.paragraphs[18].runs[8].font.name = 'Times New Roman'
            doc.paragraphs[12].text = todaysdate.strftime("%B %d, %Y")
            doc.paragraphs[12].runs[0].font.name = 'Times New Roman'
            doc.paragraphs[13].text = Residents
            doc.paragraphs[13].runs[0].font.name = 'Times New Roman'
            doc.paragraphs[14].text = StreetAddress
            doc.paragraphs[14].runs[0].font.name = 'Times New Roman'
            doc.paragraphs[15].text = CityStateZip
            doc.paragraphs[15].runs[0].font.name = 'Times New Roman'
            doc.paragraphs[27].runs[3].text = '250'
            doc.paragraphs[27].runs[3].font.name = 'Times New Roman'
            doc.save('News/{} - {} - Invoice.docx'.format(Borrower, LoanNumber))
            if values['PDF'] == True:
                word = comtypes.client.CreateObject('Word.Application')
                tempdoc = word.Documents.Open(os.path.abspath('News/{} - {} - Invoice.docx'.format(Borrower, LoanNumber)))
                tempdoc.SaveAs(os.path.abspath('News/{} - {} - Invoice.pdf'.format(Borrower, LoanNumber)), FileFormat=wdFormatPDF)
                tempdoc.Close()
                word.Quit() 
            # Coop30DayNotice
            CoopNotice.paragraphs[9].runs[2].text = Contact
            CoopNotice.paragraphs[9].runs[2].font.name = 'Times New Roman'
            CoopNotice.paragraphs[11].runs[7].text = ContactEmail
            CoopNotice.paragraphs[11].runs[7].font.name = 'Times New Roman'
            CoopNotice.paragraphs[15].runs[2].text = todaysdate.strftime("%m/%d/%Y")
            CoopNotice.paragraphs[15].runs[2].font.name = 'Times New Roman'
            CoopNotice.paragraphs[17].runs[2].text = Borrower
            CoopNotice.paragraphs[17].runs[2].font.name = 'Times New Roman'
            CoopNotice.paragraphs[19].runs[2].text = LoanNumber
            CoopNotice.paragraphs[19].runs[2].font.name = 'Times New Roman'
            CoopNotice.paragraphs[21].runs[2].text = ''
            CoopNotice.paragraphs[21].runs[2].font.name = 'Times New Roman'
            CoopNotice.save('News/{} - {} - Coop30DayNotice.docx'.format(Borrower, LoanNumber))
            if values['PDF'] == True:
                word = comtypes.client.CreateObject('Word.Application')
                tempdoc = word.Documents.Open(os.path.abspath('News/{} - {} - Coop30DayNotice.docx'.format(Borrower, LoanNumber)))
                tempdoc.SaveAs(os.path.abspath('News/{} - {} - Coop30DayNotice.pdf'.format(Borrower, LoanNumber)), FileFormat=wdFormatPDF)
                tempdoc.Close()
                word.Quit()
            # Coop Scheduling Form
            CoopScheduling.paragraphs[9].runs[2].text = Contact
            CoopScheduling.paragraphs[9].runs[2].font.name = 'Times New Roman'
            CoopScheduling.paragraphs[10].runs[6].text = ContactEmail
            CoopScheduling.paragraphs[10].runs[6].font.name = 'Times New Roman'
            CoopScheduling.paragraphs[13].runs[4].text = Borrower
            CoopScheduling.paragraphs[13].runs[4].font.name = 'Times New Roman'
            CoopScheduling.paragraphs[13].runs[11].text = LoanNumber
            CoopScheduling.paragraphs[13].runs[11].font.name = 'Times New Roman'
            CoopScheduling.save('Scheduling Forms/{} - {} - Scheduling Form.docx'.format(Borrower, LoanNumber))
            # Update Sheet
            CoopUpdateSheet.paragraphs[0].runs[0].text = Borrower
            CoopUpdateSheet.paragraphs[0].runs[0].font.name = 'Times New Roman'
            CoopUpdateSheet.paragraphs[0].runs[4].text = LoanNumber
            CoopUpdateSheet.paragraphs[0].runs[4].font.name = 'Times New Roman'
            CoopUpdateSheet.paragraphs[14].runs[3].text = Contact
            CoopUpdateSheet.paragraphs[14].runs[3].font.name = 'Times New Roman'
            CoopUpdateSheet.paragraphs[15].runs[1].text = ContactEmail
            CoopUpdateSheet.paragraphs[15].runs[1].font.name = 'Times New Roman'
            CoopUpdateSheet.paragraphs[11].runs[3].text = ''
            CoopUpdateSheet.paragraphs[11].runs[3].font.name = 'Times New Roman'
            CoopUpdateSheet.paragraphs[16].runs[1].text = PhoneNumber
            CoopUpdateSheet.paragraphs[16].runs[1].font.name = 'Times New Roman'
            CoopUpdateSheet.save('News/{} - {} - UpdateSheet.docx'.format(Borrower, LoanNumber))
            if values['PDF'] == True:
                word = comtypes.client.CreateObject('Word.Application')
                tempdoc = word.Documents.Open(os.path.abspath('News/{} - {} - UpdateSheet.docx'.format(Borrower, LoanNumber)))
                tempdoc.SaveAs(os.path.abspath('News/{} - {} - UpdateSheet.pdf'.format(Borrower, LoanNumber)), FileFormat=wdFormatPDF)
                tempdoc.Close()
                word.Quit()
            # Coop Email with PDF Attachments
            if values['PDF'] == True and values['Notice'] == True:
                message = outlook.CreateItem(0)
                message.To = ContactEmail
                # message.CC = 'chrisheyer0@gmail.com'
                message.Subject = '{} - {} - Initial Notice'.format(Borrower, LoanNumber)
                message.HTMLBody = CoopEmail
                message.Attachments.Add('{}/News/{} - {} - Invoice.pdf'.format(CurrentDirectory, Borrower, LoanNumber))
                message.Attachments.Add('{}/News/{} - {} - Coop30DayNotice.pdf'.format(CurrentDirectory, Borrower, LoanNumber))
                message.Send()
                window['Output'].update('Coop Initial Notice sent to: {} - {}'.format(ContactEmail, datetime.datetime.now().strftime("%m/%d/%Y at %I:%M:%S %p")))
            # Coop Email with DOCX Attachments
            if values['PDF'] == False and values['Notice'] == True:
                message = outlook.CreateItem(0)
                message.To = ContactEmail
                # message.CC = 'chrisheyer0@gmail.com'
                message.Subject = '{} - {} - Initial Notice'.format(Borrower, LoanNumber)
                message.HTMLBody = CoopEmail
                message.Attachments.Add('{}/News/{} - {} - Invoice.docx'.format(CurrentDirectory, Borrower, LoanNumber))
                message.Attachments.Add('{}/News/{} - {} - Coop30DayNotice.docx'.format(CurrentDirectory, Borrower, LoanNumber))
                message.Send()
                window['Output'].update('Coop Initial Notice sent to: {} - {}'.format(ContactEmail, datetime.datetime.now().strftime("%m/%d/%Y at %I:%M:%S %p")))


        # Coop 1st + HELOC
        elif values['LoanType'] == 'Coop 1st + HELOC':
            # Appending to Excel
            newRowLocation = ws.max_row + 1
            ws.cell(column = 1, row = newRowLocation, value = Borrower)
            ws.cell(column = 2, row = newRowLocation, value = LoanNumber)
            ws.cell(column = 5, row = newRowLocation, value = Contact)
            ws.cell(column = 6, row = newRowLocation, value = ContactEmail)
            ws.cell(column = 3, row = newRowLocation, value = todaysdate.strftime("%m/%d/%Y"))
            wb.save('NEWS_Status.xlsx')
            wb.close()
            # Invoice
            doc.paragraphs[18].runs[2].text = Borrower
            doc.paragraphs[18].runs[2].font.name = 'Times New Roman'
            doc.paragraphs[18].runs[8].text = '{} & {}'.format(LoanNumber, HelocNumber)
            doc.paragraphs[18].runs[8].font.name = 'Times New Roman'
            doc.paragraphs[12].text = todaysdate.strftime("%B %d, %Y")
            doc.paragraphs[12].runs[0].font.name = 'Times New Roman'
            doc.paragraphs[13].text = Residents
            doc.paragraphs[13].runs[0].font.name = 'Times New Roman'
            doc.paragraphs[14].text = StreetAddress
            doc.paragraphs[14].runs[0].font.name = 'Times New Roman'
            doc.paragraphs[15].text = CityStateZip
            doc.paragraphs[15].runs[0].font.name = 'Times New Roman'
            doc.paragraphs[27].runs[3].text = '400'
            doc.paragraphs[27].runs[3].font.name = 'Times New Roman'
            doc.save('News/{} - {} - Invoice.docx'.format(Borrower, LoanNumber))
            if values['PDF'] == True:
                word = comtypes.client.CreateObject('Word.Application')
                tempdoc = word.Documents.Open(os.path.abspath('News/{} - {} - Invoice.docx'.format(Borrower, LoanNumber)))
                tempdoc.SaveAs(os.path.abspath('News/{} - {} - Invoice.pdf'.format(Borrower, LoanNumber)), FileFormat=wdFormatPDF)
                tempdoc.Close()
                word.Quit()
            # 30 Day Notice
            Coop1stHELOCNotice.paragraphs[9].runs[2].text = Contact
            Coop1stHELOCNotice.paragraphs[9].runs[2].font.name = 'Times New Roman'
            Coop1stHELOCNotice.paragraphs[11].runs[6].text = ContactEmail
            Coop1stHELOCNotice.paragraphs[11].runs[6].font.name = 'Times New Roman'
            Coop1stHELOCNotice.paragraphs[15].runs[2].text = todaysdate.strftime("%m/%d/%Y")
            Coop1stHELOCNotice.paragraphs[15].runs[2].font.name = 'Times New Roman'
            Coop1stHELOCNotice.paragraphs[17].runs[2].text = Borrower
            Coop1stHELOCNotice.paragraphs[17].runs[2].font.name = 'Times New Roman'
            Coop1stHELOCNotice.paragraphs[19].runs[2].text = LoanNumber
            Coop1stHELOCNotice.paragraphs[19].runs[2].font.name = 'Times New Roman'
            Coop1stHELOCNotice.paragraphs[21].runs[2].text = HelocNumber
            Coop1stHELOCNotice.paragraphs[21].runs[2].font.name = 'Times New Roman'
            Coop1stHELOCNotice.save('News/{} - {} - Coop30DayNotice.docx'.format(Borrower, LoanNumber))
            if values['PDF'] == True:
                word = comtypes.client.CreateObject('Word.Application')
                tempdoc = word.Documents.Open(os.path.abspath('News/{} - {} - Coop30DayNotice.docx'.format(Borrower, LoanNumber)))
                tempdoc.SaveAs(os.path.abspath('News/{} - {} - Coop30DayNotice.pdf'.format(Borrower, LoanNumber)), FileFormat=wdFormatPDF)
                tempdoc.Close()
                word.Quit()
            # Coop Scheduling Form
            CoopScheduling.paragraphs[9].runs[2].text = Contact
            CoopScheduling.paragraphs[9].runs[2].font.name = 'Times New Roman'
            CoopScheduling.paragraphs[10].runs[6].text = ContactEmail
            CoopScheduling.paragraphs[10].runs[6].font.name = 'Times New Roman'
            CoopScheduling.paragraphs[13].runs[4].text = Borrower
            CoopScheduling.paragraphs[13].runs[4].font.name = 'Times New Roman'
            CoopScheduling.paragraphs[13].runs[11].text = '{} & {}'.format(LoanNumber, HelocNumber)
            CoopScheduling.paragraphs[13].runs[11].font.name = 'Times New Roman'
            CoopScheduling.save('Scheduling Forms/{} - {} - Scheduling Form.docx'.format(Borrower, LoanNumber))
            # Update Sheet
            CoopUpdateSheet.paragraphs[0].runs[0].text = Borrower
            CoopUpdateSheet.paragraphs[0].runs[0].font.name = 'Times New Roman'
            CoopUpdateSheet.paragraphs[0].runs[4].text = LoanNumber
            CoopUpdateSheet.paragraphs[0].runs[4].font.name = 'Times New Roman'
            CoopUpdateSheet.paragraphs[14].runs[3].text = Contact
            CoopUpdateSheet.paragraphs[14].runs[3].font.name = 'Times New Roman'
            CoopUpdateSheet.paragraphs[15].runs[1].text = ContactEmail
            CoopUpdateSheet.paragraphs[15].runs[1].font.name = 'Times New Roman'
            CoopUpdateSheet.paragraphs[11].runs[3].text = HelocNumber
            CoopUpdateSheet.paragraphs[11].runs[3].font.name = 'Times New Roman'
            CoopUpdateSheet.paragraphs[16].runs[1].text = PhoneNumber
            CoopUpdateSheet.paragraphs[16].runs[1].font.name = 'Times New Roman'
            CoopUpdateSheet.save('News/{} - {} - UpdateSheet.docx'.format(Borrower, LoanNumber))
            if values['PDF'] == True:
                word = comtypes.client.CreateObject('Word.Application')
                tempdoc = word.Documents.Open(os.path.abspath('News/{} - {} - UpdateSheet.docx'.format(Borrower, LoanNumber)))
                tempdoc.SaveAs(os.path.abspath('News/{} - {} - UpdateSheet.pdf'.format(Borrower, LoanNumber)), FileFormat=wdFormatPDF)
                tempdoc.Close()
                word.Quit()
            # Coop Email with PDF Attachments
            if values['PDF'] == True and values['Notice'] == True:
                message = outlook.CreateItem(0)
                message.To = ContactEmail
                # message.CC = 'chrisheyer0@gmail.com'
                message.Subject = '{} - {} - Initial Notice'.format(Borrower, LoanNumber)
                message.HTMLBody = CoopEmail
                message.Attachments.Add('{}/News/{} - {} - Invoice.pdf'.format(CurrentDirectory, Borrower, LoanNumber))
                message.Attachments.Add('{}/News/{} - {} - Coop30DayNotice.pdf'.format(CurrentDirectory, Borrower, LoanNumber))
                message.Send()
                window['Output'].update('Coop Initial Notice sent to: {} - {}'.format(ContactEmail, datetime.datetime.now().strftime("%m/%d/%Y at %I:%M:%S %p")))
            # Coop Email with DOCX Attachments
            if values['PDF'] == False and values['Notice'] == True:
                message = outlook.CreateItem(0)
                message.To = ContactEmail
                # message.CC = 'chrisheyer0@gmail.com'
                message.Subject = '{} - {} - Initial Notice'.format(Borrower, LoanNumber)
                message.HTMLBody = CoopEmail
                message.Attachments.Add('{}/News/{} - {} - Invoice.docx'.format(CurrentDirectory, Borrower, LoanNumber))
                message.Attachments.Add('{}/News/{} - {} - Coop30DayNotice.docx'.format(CurrentDirectory, Borrower, LoanNumber))
                message.Send()
                window['Output'].update('Coop 1st + HELOC Initial Notice sent to: {} - {}'.format(ContactEmail, datetime.datetime.now().strftime("%m/%d/%Y at %I:%M:%S %p")))


        # Coop HELOC
        elif values['LoanType'] == 'Coop HELOC':
            LoanNumber = HelocNumber
            # Appending to Excel
            newRowLocation = ws.max_row + 1
            ws.cell(column = 1, row = newRowLocation, value = Borrower)
            ws.cell(column = 2, row = newRowLocation, value = LoanNumber)
            ws.cell(column = 5, row = newRowLocation, value = Contact)
            ws.cell(column = 6, row = newRowLocation, value = ContactEmail)
            ws.cell(column = 3, row = newRowLocation, value = todaysdate.strftime("%m/%d/%Y"))
            wb.save('NEWS_Status.xlsx')
            wb.close()
            # Invoice
            doc.paragraphs[18].runs[2].text = Borrower
            doc.paragraphs[18].runs[2].font.name = 'Times New Roman'
            doc.paragraphs[18].runs[8].text = LoanNumber
            doc.paragraphs[18].runs[8].font.name = 'Times New Roman'
            doc.paragraphs[12].text = todaysdate.strftime("%B %d, %Y")
            doc.paragraphs[12].runs[0].font.name = 'Times New Roman'
            doc.paragraphs[13].text = Residents
            doc.paragraphs[13].runs[0].font.name = 'Times New Roman'
            doc.paragraphs[14].text = StreetAddress
            doc.paragraphs[14].runs[0].font.name = 'Times New Roman'
            doc.paragraphs[15].text = CityStateZip
            doc.paragraphs[15].runs[0].font.name = 'Times New Roman'
            doc.paragraphs[27].runs[3].text = '375'
            doc.paragraphs[27].runs[3].font.name = 'Times New Roman'
            doc.save('News/{} - {} - Invoice.docx'.format(Borrower, LoanNumber))
            if values['PDF'] == True:
                word = comtypes.client.CreateObject('Word.Application')
                tempdoc = word.Documents.Open(os.path.abspath('News/{} - {} - Invoice.docx'.format(Borrower, LoanNumber)))
                tempdoc.SaveAs(os.path.abspath('News/{} - {} - Invoice.pdf'.format(Borrower, LoanNumber)), FileFormat=wdFormatPDF)
                tempdoc.Close()
                word.Quit()
            # 30 Day Notice
            CoopHELOCNotice.paragraphs[9].runs[2].text = Contact
            CoopHELOCNotice.paragraphs[9].runs[2].font.name = 'Times New Roman'
            CoopHELOCNotice.paragraphs[11].runs[6].text = ContactEmail
            CoopHELOCNotice.paragraphs[11].runs[6].font.name = 'Times New Roman'
            CoopHELOCNotice.paragraphs[15].runs[2].text = todaysdate.strftime("%m/%d/%Y")
            CoopHELOCNotice.paragraphs[15].runs[2].font.name = 'Times New Roman'
            CoopHELOCNotice.paragraphs[17].runs[2].text = Borrower
            CoopHELOCNotice.paragraphs[17].runs[2].font.name = 'Times New Roman'
            CoopHELOCNotice.paragraphs[19].runs[2].text = ''
            CoopHELOCNotice.paragraphs[19].runs[2].font.name = 'Times New Roman'
            CoopHELOCNotice.paragraphs[21].runs[2].text = LoanNumber
            CoopHELOCNotice.paragraphs[21].runs[2].font.name = 'Times New Roman'
            CoopHELOCNotice.save('News/{} - {} - Coop30DayNotice.docx'.format(Borrower, LoanNumber))
            if values['PDF'] == True:
                word = comtypes.client.CreateObject('Word.Application')
                tempdoc = word.Documents.Open(os.path.abspath('News/{} - {} - Coop30DayNotice.docx'.format(Borrower, LoanNumber)))
                tempdoc.SaveAs(os.path.abspath('News/{} - {} - Coop30DayNotice.pdf'.format(Borrower, LoanNumber)), FileFormat=wdFormatPDF)
                tempdoc.Close()
                word.Quit()
            # Coop Scheduling Form
            CoopScheduling.paragraphs[9].runs[2].text = Contact
            CoopScheduling.paragraphs[9].runs[2].font.name = 'Times New Roman'
            CoopScheduling.paragraphs[10].runs[6].text = ContactEmail
            CoopScheduling.paragraphs[10].runs[6].font.name = 'Times New Roman'
            CoopScheduling.paragraphs[13].runs[4].text = Borrower
            CoopScheduling.paragraphs[13].runs[4].font.name = 'Times New Roman'
            CoopScheduling.paragraphs[13].runs[11].text = LoanNumber
            CoopScheduling.paragraphs[13].runs[11].font.name = 'Times New Roman'
            CoopScheduling.save('Scheduling Forms/{} - {} - Scheduling Form.docx'.format(Borrower, LoanNumber))
            # Update Sheet
            CoopUpdateSheet.paragraphs[0].runs[0].text = Borrower
            CoopUpdateSheet.paragraphs[0].runs[0].font.name = 'Times New Roman'
            CoopUpdateSheet.paragraphs[0].runs[4].text = LoanNumber
            CoopUpdateSheet.paragraphs[0].runs[4].font.name = 'Times New Roman'
            CoopUpdateSheet.paragraphs[14].runs[3].text = Contact
            CoopUpdateSheet.paragraphs[14].runs[3].font.name = 'Times New Roman'
            CoopUpdateSheet.paragraphs[15].runs[1].text = ContactEmail
            CoopUpdateSheet.paragraphs[15].runs[1].font.name = 'Times New Roman'
            CoopUpdateSheet.paragraphs[11].runs[3].text = ''
            CoopUpdateSheet.paragraphs[11].runs[3].font.name = 'Times New Roman'
            CoopUpdateSheet.paragraphs[16].runs[1].text = PhoneNumber
            CoopUpdateSheet.paragraphs[16].runs[1].font.name = 'Times New Roman'
            CoopUpdateSheet.save('News/{} - {} - UpdateSheet.docx'.format(Borrower, LoanNumber))
            if values['PDF'] == True:
                word = comtypes.client.CreateObject('Word.Application')
                tempdoc = word.Documents.Open(os.path.abspath('News/{} - {} - UpdateSheet.docx'.format(Borrower, LoanNumber)))
                tempdoc.SaveAs(os.path.abspath('News/{} - {} - UpdateSheet.pdf'.format(Borrower, LoanNumber)), FileFormat=wdFormatPDF)
                tempdoc.Close()
                word.Quit()
            # Coop Email with PDF Attachments
            if values['PDF'] == True and values['Notice'] == True:
                message = outlook.CreateItem(0)
                message.To = ContactEmail
                # message.CC = 'chrisheyer0@gmail.com'
                message.Subject = '{} - {} - Initial Notice'.format(Borrower, LoanNumber)
                message.HTMLBody = CoopEmail
                message.Attachments.Add('{}/News/{} - {} - Invoice.pdf'.format(CurrentDirectory, Borrower, LoanNumber))
                message.Attachments.Add('{}/News/{} - {} - Coop30DayNotice.pdf'.format(CurrentDirectory, Borrower, LoanNumber))
                message.Send()
                window['Output'].update('Coop Initial Notice sent to: {} - {}'.format(ContactEmail, datetime.datetime.now().strftime("%m/%d/%Y at %I:%M:%S %p")))
            # Coop Email with DOCX Attachments
            if values['PDF'] == False and values['Notice'] == True:
                message = outlook.CreateItem(0)
                message.To = ContactEmail
                # message.CC = 'chrisheyer0@gmail.com'
                message.Subject = '{} - {} - Initial Notice'.format(Borrower, LoanNumber)
                message.HTMLBody = CoopEmail
                message.Attachments.Add('{}/News/{} - {} - Invoice.docx'.format(CurrentDirectory, Borrower, LoanNumber))
                message.Attachments.Add('{}/News/{} - {} - Coop30DayNotice.docx'.format(CurrentDirectory, Borrower, LoanNumber))
                message.Send()
                window['Output'].update('Coop HELOC Initial Notice sent to: {} - {}'.format(ContactEmail, datetime.datetime.now().strftime("%m/%d/%Y at %I:%M:%S %p")))




    # Clearing Form Data     
    if event == 'Clear':
        window['-borrower-'].update('')
        window['-loannumber-'].update('')
        window['-helocnumber-'].update('')
        window['-contact-'].update('')
        window['-contactemail-'].update('')
        window['-phonenumber-'].update('')
        window['-residents-'].update('')
        window['-streetaddress-'].update('')
        window['-citystatezip-'].update('')
        window['Output'].update('Fields Cleared')

window.close()

