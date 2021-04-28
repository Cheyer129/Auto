import docx
import datetime
from openpyxl import load_workbook

# Opening up the original forms to modify
Invoice = docx.Document('Toscano Files/InvoiceSheet.docx')
CoopNotice = docx.Document('Toscano Files/Coop30DayNotice.docx')
CemaNotice = docx.Document('Toscano Files/CEMA30DayNotice.docx')
CoopScheduling = docx.Document('Toscano Files/CoopSchedulingForm.docx')
CemaScheduling = docx.Document('Toscano Files/CEMASchedulingForm.docx')
CoopUpdateSheet = docx.Document('Toscano Files/CoopUpdateSheet.docx')
CemaUpdateSheet = docx.Document('Toscano Files/CEMAUpdateSheet.docx')
CemaAssignmentForm = docx.Document('Toscano Files/AssignmentForm.docx')
CemaHELOCNotice = docx.Document('Toscano Files/HELOCCEMA30DayNotice.docx')
Coop1stHELOCNotice = docx.Document('Toscano Files/1st+HELOCCoop30DayNotice.docx')
CoopHELOCNotice = docx.Document('Toscano Files/HELOCCoop30DayNotice.docx')

# Accessing the excel file
wb = load_workbook('NEWS_Status.xlsx')
ws = wb['Sheet1']

def MakeInvoice(Borrower, LoanNumber, Residents, StreetAddress, CityStateZip, Fee):
        Invoice.paragraphs[18].runs[2].text = Borrower
        Invoice.paragraphs[18].runs[2].font.name = 'Times New Roman'
        Invoice.paragraphs[18].runs[8].text = LoanNumber
        Invoice.paragraphs[18].runs[8].font.name = 'Times New Roman'
        Invoice.paragraphs[12].text = datetime.datetime.now().strftime("%B %d, %Y")
        Invoice.paragraphs[12].runs[0].font.name = 'Times New Roman'
        Invoice.paragraphs[13].text = Residents
        Invoice.paragraphs[13].runs[0].font.name = 'Times New Roman'
        Invoice.paragraphs[14].text = StreetAddress
        Invoice.paragraphs[14].runs[0].font.name = 'Times New Roman'
        Invoice.paragraphs[15].text = CityStateZip
        Invoice.paragraphs[15].runs[0].font.name = 'Times New Roman'
        Invoice.paragraphs[27].runs[3].text = Fee
        Invoice.paragraphs[27].runs[3].font.name = 'Times New Roman'

def AppendToExcelCEMA(Borrower, LoanNumber, Contact, ContactEmail):
    newRowLocation = ws.max_row + 1
    ws.cell(column = 1, row = newRowLocation, value = Borrower)
    ws.cell(column = 2, row = newRowLocation, value = LoanNumber)
    ws.cell(column = 5, row = newRowLocation, value = Contact)
    ws.cell(column = 6, row = newRowLocation, value = ContactEmail)
    ws.cell(column = 3, row = newRowLocation, value = datetime.datetime.now().strftime("%m/%d/%Y"))
    wb.save('NEWS_Status.xlsx')
    wb.close()

def AppendToExcelCoop(Borrower, LoanNumber, Contact, ContactEmail):
    newRowLocation = ws.max_row + 1
    ws.cell(column = 1, row = newRowLocation, value = Borrower)
    ws.cell(column = 2, row = newRowLocation, value = LoanNumber)
    ws.cell(column = 5, row = newRowLocation, value = Contact)
    ws.cell(column = 6, row = newRowLocation, value = ContactEmail)
    ws.cell(column = 3, row = newRowLocation, value = datetime.datetime.now().strftime("%m/%d/%Y"))
    wb.save('NEWS_Status.xlsx')
    wb.close()
