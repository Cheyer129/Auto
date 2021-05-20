import PySimpleGUI as sg
import docx
import datetime
import time
from docx2pdf import convert
import os
import win32com.client as client
from pathlib import Path
import openpyxl
from openpyxl import load_workbook
import comtypes.client

# Creating Path
CurrentDirectory = Path.cwd()

# Creating date object
todaysdate = datetime.datetime.now()

# Creating a list of users
users = os.listdir('Toscano Files/USERS')

# Accessing the excel file
wb = load_workbook('NEWS_Status.xlsx')
ws = wb['Sheet1']

# Accessing Word Application
word = comtypes.client.CreateObject('Word.Application')

# Emails HTMLs
CemaEmail = open('Toscano Files/CemaEmail.html', 'r').read()
CoopEmail = open('Toscano Files/CoopEmail.html', 'r').read()

# Accessing Outlook Application
outlook = client.Dispatch("Outlook.Application")

# Opening up the Original Invoice Sheet to modify
Invoice = docx.Document('Toscano Files/InvoiceSheet.docx')
CoopNotice = docx.Document('Toscano Files/Coop30DayNotice.docx')
CemaNotice = docx.Document('Toscano Files/CEMA30DayNotice.docx')
CoopScheduling = docx.Document('Toscano Files/CoopSchedulingForm.docx')
CemaScheduling = docx.Document('Toscano Files/CEMASchedulingForm.docx')
CoopUpdateSheet = docx.Document('Toscano Files/CoopUpdateSheet.docx')
CemaUpdateSheet = docx.Document('Toscano Files/CEMAUpdateSheet.docx')
CoopHELOCNotice = docx.Document('Toscano Files/HELOCCoop30DayNotice.docx')
Coop1stHELOCNotice = docx.Document('Toscano Files/1st+HELOCCoop30DayNotice.docx')
CemaHELOCNotice = docx.Document('Toscano Files/HELOCCEMA30DayNotice.docx')

# Theme COLORS!
sg.theme('GreenMono')

# All the stuff inside your window.
layout = [  [sg.Text('New File Form Generator', font = ('Calibri', 22), pad = ((5,5),(5,20)))],
            [sg.Combo(users, key = 'user', pad = ((5,5),(5,15)))],
            [sg.Text('CC', pad = ((5,5),(5,5))), sg.InputText(key = 'cc')],
            [sg.Text('Borrower', pad = ((5,43),(5,5))), sg.InputText(key = '-borrower-')],
            [sg.Text('Loan Number', pad = ((5,16),(5,5))), sg.InputText(key = '-loannumber-')],
            [sg.Text('HELOC #', pad = ((5,38),(5,5))),sg.InputText(key = '-helocnumber-')],
            [sg.Text('Contact', pad = ((5,48),(5,5))), sg.InputText(key = '-contact-')],
            [sg.Text('Contact Email', pad = ((5,11),(5,5))), sg.InputText(key = '-contactemail-')],
            [sg.Text('Phone Number', pad = ((5,7),(5,5))), sg.InputText(key = '-phonenumber-')],
            [sg.Text('Residents', pad = ((5,34),(5,5))), sg.InputText(key = '-residents-')],
            [sg.Text('Street Address', pad = ((5,5),(5,5))), sg.InputText(key = '-streetaddress-')],
            [sg.Text('City, State Zip', pad = ((5,9),(5,20))), sg.InputText(key = '-citystatezip-', pad = ((5,5),(5,20)))],
            [sg.Combo(['CEMA', 'CEMA HELOC', 'Coop', 'Coop 1st + HELOC', 'Coop HELOC'], key = 'LoanType', pad = ((5,5),(5,20)))],
            [sg.Checkbox('Create PDFs', key = 'PDF', pad = ((5,5,),(5,12))), sg.Checkbox('Send Notice', key = 'Notice', pad = ((5,5,),(5,12)))],
            [sg.Button('Enter'), sg.Button('Clear'), sg.Button('Exit')],
            [sg.Text('', size = (55,1), auto_size_text = False, pad = ((0,5),(5,0)), key = 'Output')],
            [sg.Text('', size = (55,1), auto_size_text = False, pad = ((0,5),(0,5)), key = 'Output2')],
            ]

# Create the Window 
window = sg.Window('News Generator', layout, default_element_size = (45,1), location = (50,150), resizable = True)


class Payoffs:
    def __init__(self, loan_number, heloc_number, borrower, loan_type, contact, 
                 contact_email, phone_number, residents, street_address, city_state_zip):
        self.loan_number = loan_number
        self.heloc_number = heloc_number
        self.borrower = borrower
        self.loan_type = loan_type
        self.contact = contact
        self.contact_email = contact_email
        self.phone_number = phone_number
        self.residents = residents
        self.street_address = street_address
        self.city_state_zip = city_state_zip

    def add_to_excel(self):
        if self.loan_type == 'CEMA':
            newRowLocation = ws.max_row + 1
            ws.cell(column = 1, row = newRowLocation, value = self.borrower)
            ws.cell(column = 2, row = newRowLocation, value = self.loan_number)
            ws.cell(column = 5, row = newRowLocation, value = self.contact)
            ws.cell(column = 6, row = newRowLocation, value = self.contact_email)
            ws.cell(column = 4, row = newRowLocation, value = todaysdate.strftime("%m/%d/%Y"))
            wb.save('NEWS_Status.xlsx')
            wb.close()
        elif self.loan_type == 'CEMA HELOC':
            newRowLocation = ws.max_row + 1
            ws.cell(column = 1, row = newRowLocation, value = self.borrower)
            ws.cell(column = 2, row = newRowLocation, value = self.heloc_number)
            ws.cell(column = 5, row = newRowLocation, value = self.contact)
            ws.cell(column = 6, row = newRowLocation, value = self.contact_email)
            ws.cell(column = 4, row = newRowLocation, value = todaysdate.strftime("%m/%d/%Y"))
            wb.save('NEWS_Status.xlsx')
            wb.close()
        elif self.loan_type == 'Coop' or self.loan_type == 'Coop 1st + HELOC':
            newRowLocation = ws.max_row + 1
            ws.cell(column = 1, row = newRowLocation, value = self.borrower)
            ws.cell(column = 2, row = newRowLocation, value = self.loan_number)
            ws.cell(column = 5, row = newRowLocation, value = self.contact)
            ws.cell(column = 6, row = newRowLocation, value = self.contact_email)
            ws.cell(column = 3, row = newRowLocation, value = todaysdate.strftime("%m/%d/%Y"))
            wb.save('NEWS_Status.xlsx')
            wb.close()
        elif self.loan_type == 'Coop HELOC':
            newRowLocation = ws.max_row + 1
            ws.cell(column = 1, row = newRowLocation, value = self.borrower)
            ws.cell(column = 2, row = newRowLocation, value = self.heloc_number)
            ws.cell(column = 5, row = newRowLocation, value = self.contact)
            ws.cell(column = 6, row = newRowLocation, value = self.contact_email)
            ws.cell(column = 3, row = newRowLocation, value = todaysdate.strftime("%m/%d/%Y"))
            wb.save('NEWS_Status.xlsx')
            wb.close()
    
    def make_invoice(self, number, fee):
        Invoice.paragraphs[18].runs[2].text = self.borrower
        Invoice.paragraphs[18].runs[2].font.name = 'Times New Roman'
        Invoice.paragraphs[18].runs[8].text = number
        Invoice.paragraphs[18].runs[8].font.name = 'Times New Roman'
        Invoice.paragraphs[12].text = datetime.datetime.now().strftime("%B %d, %Y")
        Invoice.paragraphs[12].runs[0].font.name = 'Times New Roman'
        Invoice.paragraphs[13].text = self.residents
        Invoice.paragraphs[13].runs[0].font.name = 'Times New Roman'
        Invoice.paragraphs[14].text = self.street_address
        Invoice.paragraphs[14].runs[0].font.name = 'Times New Roman'
        Invoice.paragraphs[15].text = self.city_state_zip
        Invoice.paragraphs[15].runs[0].font.name = 'Times New Roman'
        Invoice.paragraphs[27].runs[3].text = str(fee)
        Invoice.paragraphs[27].runs[3].font.name = 'Times New Roman'
        if self.loan_type == 'Coop 1st + HELOC':
            Invoice.save('News/{} - {} - Invoice.docx'.format(self.borrower, self.loan_number))
        else:
            Invoice.save('News/{} - {} - Invoice.docx'.format(self.borrower, number))

    def make_update_sheet(self, number):
        if self.loan_type == 'CEMA' or self.loan_type == 'CEMA HELOC':
            CemaUpdateSheet.paragraphs[0].runs[0].text = self.borrower
            CemaUpdateSheet.paragraphs[0].runs[0].font.name = 'Times New Roman'
            CemaUpdateSheet.paragraphs[0].runs[5].text = number
            CemaUpdateSheet.paragraphs[0].runs[5].font.name = 'Times New Roman'
            CemaUpdateSheet.paragraphs[16].runs[2].text = self.contact
            CemaUpdateSheet.paragraphs[16].runs[2].font.name = 'Times New Roman'
            CemaUpdateSheet.paragraphs[17].runs[1].text = self.contact_email
            CemaUpdateSheet.paragraphs[17].runs[1].font.name = 'Times New Roman'
            CemaUpdateSheet.paragraphs[18].runs[1].text = self.phone_number
            CemaUpdateSheet.paragraphs[18].runs[1].font.name = 'Times New Roman'
            CemaUpdateSheet.save('News/{} - {} - UpdateSheet.docx'.format(self.borrower, number))
        elif self.loan_type == 'Coop' or self.loan_type == 'Coop HELOC':
            CoopUpdateSheet.paragraphs[0].runs[0].text = self.borrower
            CoopUpdateSheet.paragraphs[0].runs[0].font.name = 'Times New Roman'
            CoopUpdateSheet.paragraphs[0].runs[4].text = number
            CoopUpdateSheet.paragraphs[0].runs[4].font.name = 'Times New Roman'
            CoopUpdateSheet.paragraphs[14].runs[3].text = self.contact
            CoopUpdateSheet.paragraphs[14].runs[3].font.name = 'Times New Roman'
            CoopUpdateSheet.paragraphs[15].runs[1].text = self.contact_email
            CoopUpdateSheet.paragraphs[15].runs[1].font.name = 'Times New Roman'
            CoopUpdateSheet.paragraphs[11].runs[3].text = ''
            CoopUpdateSheet.paragraphs[11].runs[3].font.name = 'Times New Roman'
            CoopUpdateSheet.paragraphs[16].runs[1].text = self.phone_number
            CoopUpdateSheet.paragraphs[16].runs[1].font.name = 'Times New Roman'
            CoopUpdateSheet.save('News/{} - {} - UpdateSheet.docx'.format(self.borrower, number))
        elif self.loan_type == 'Coop 1st + HELOC':
            CoopUpdateSheet.paragraphs[0].runs[0].text = self.borrower
            CoopUpdateSheet.paragraphs[0].runs[0].font.name = 'Times New Roman'
            CoopUpdateSheet.paragraphs[0].runs[4].text = number
            CoopUpdateSheet.paragraphs[0].runs[4].font.name = 'Times New Roman'
            CoopUpdateSheet.paragraphs[14].runs[3].text = self.contact
            CoopUpdateSheet.paragraphs[14].runs[3].font.name = 'Times New Roman'
            CoopUpdateSheet.paragraphs[15].runs[1].text = self.contact_email
            CoopUpdateSheet.paragraphs[15].runs[1].font.name = 'Times New Roman'
            CoopUpdateSheet.paragraphs[11].runs[3].text = self.heloc_number
            CoopUpdateSheet.paragraphs[11].runs[3].font.name = 'Times New Roman'
            CoopUpdateSheet.paragraphs[16].runs[1].text = self.phone_number
            CoopUpdateSheet.paragraphs[16].runs[1].font.name = 'Times New Roman'
            CoopUpdateSheet.save('News/{} - {} - UpdateSheet.docx'.format(self.borrower, number))
    
    def make_scheduling_form(self, number):
        if self.loan_type == 'CEMA' or self.loan_type == 'CEMA HELOC':
            CemaScheduling.paragraphs[9].runs[2].text = self.contact
            CemaScheduling.paragraphs[9].runs[2].font.name = 'Times New Roman'
            CemaScheduling.paragraphs[10].runs[8].text = self.contact_email
            CemaScheduling.paragraphs[10].runs[8].font.name = 'Times New Roman'
            CemaScheduling.paragraphs[13].runs[5].text = self.borrower
            CemaScheduling.paragraphs[13].runs[5].font.name = 'Times New Roman'
            CemaScheduling.paragraphs[13].runs[11].text = number
            CemaScheduling.paragraphs[13].runs[11].font.name = 'Times New Roman'
            CemaScheduling.save('Scheduling Forms/{} - {} - Scheduling Form.docx'.format(self.borrower, number))
        if self.loan_type == 'Coop' or self.loan_type == 'Coop HELOC' or self.loan_type == 'Coop 1st + HELOC':
            CoopScheduling.paragraphs[9].runs[2].text = self.contact
            CoopScheduling.paragraphs[9].runs[2].font.name = 'Times New Roman'
            CoopScheduling.paragraphs[10].runs[7].text = self.contact_email
            CoopScheduling.paragraphs[10].runs[7].font.name = 'Times New Roman'
            CoopScheduling.paragraphs[13].runs[4].text = self.borrower
            CoopScheduling.paragraphs[13].runs[4].font.name = 'Times New Roman'
            if self.loan_type == 'Coop':
                CoopScheduling.paragraphs[13].runs[11].text = self.loan_number
                CoopScheduling.paragraphs[13].runs[11].font.name = 'Times New Roman'
                CoopScheduling.save('Scheduling Forms/{} - {} - Scheduling Form.docx'.format(self.borrower, number))
            elif self.loan_type == 'Coop HELOC':
                CoopScheduling.paragraphs[13].runs[11].text = self.heloc_number
                CoopScheduling.paragraphs[13].runs[11].font.name = 'Times New Roman'
                CoopScheduling.save('Scheduling Forms/{} - {} - Scheduling Form.docx'.format(self.borrower, number))
            elif self.loan_type == 'Coop 1st + HELOC':
                CoopScheduling.paragraphs[13].runs[11].text = '{} & {}'.format(self.loan_number, self.heloc_number)
                CoopScheduling.paragraphs[13].runs[11].font.name = 'Times New Roman'
                CoopScheduling.save('Scheduling Forms/{} - {} - Scheduling Form.docx'.format(self.borrower, self.loan_number))

    def make_initial_notice(self):
        if self.loan_type == 'CEMA':
            CemaNotice.paragraphs[9].runs[2].text = self.contact
            CemaNotice.paragraphs[9].runs[2].font.name = 'Times New Roman'
            CemaNotice.paragraphs[11].runs[5].text = self.contact_email
            CemaNotice.paragraphs[11].runs[5].font.name = 'Times New Roman'
            CemaNotice.paragraphs[15].runs[2].text = todaysdate.strftime("%m/%d/%Y")
            CemaNotice.paragraphs[15].runs[2].font.name = 'Times New Roman'
            CemaNotice.paragraphs[17].runs[3].text = self.borrower
            CemaNotice.paragraphs[17].runs[3].font.name = 'Times New Roman'
            CemaNotice.paragraphs[17].runs[8].text = self.loan_number
            CemaNotice.paragraphs[17].runs[8].font.name = 'Times New Roman'
            CemaNotice.save('News/{} - {} - CEMA30DayNotice.docx'.format(self.borrower, self.loan_number))
        elif self.loan_type == 'CEMA HELOC':
            CemaHELOCNotice.paragraphs[9].runs[2].text = self.contact
            CemaHELOCNotice.paragraphs[9].runs[2].font.name = 'Times New Roman'
            CemaHELOCNotice.paragraphs[11].runs[6].text = self.contact_email
            CemaHELOCNotice.paragraphs[11].runs[6].font.name = 'Times New Roman'
            CemaHELOCNotice.paragraphs[15].runs[2].text = todaysdate.strftime("%m/%d/%Y")
            CemaHELOCNotice.paragraphs[15].runs[2].font.name = 'Times New Roman'
            CemaHELOCNotice.paragraphs[17].runs[3].text = self.borrower
            CemaHELOCNotice.paragraphs[17].runs[3].font.name = 'Times New Roman'
            CemaHELOCNotice.paragraphs[17].runs[8].text = self.heloc_number
            CemaHELOCNotice.paragraphs[17].runs[8].font.name = 'Times New Roman'
            CemaHELOCNotice.save('News/{} - {} - CEMA30DayNotice.docx'.format(self.borrower, self.heloc_number))
        elif self.loan_type == 'Coop':
            CoopNotice.paragraphs[9].runs[2].text = self.contact
            CoopNotice.paragraphs[9].runs[2].font.name = 'Times New Roman'
            CoopNotice.paragraphs[11].runs[7].text = self.contact_email
            CoopNotice.paragraphs[11].runs[7].font.name = 'Times New Roman'
            CoopNotice.paragraphs[15].runs[2].text = todaysdate.strftime("%m/%d/%Y")
            CoopNotice.paragraphs[15].runs[2].font.name = 'Times New Roman'
            CoopNotice.paragraphs[17].runs[2].text = self.borrower
            CoopNotice.paragraphs[17].runs[2].font.name = 'Times New Roman'
            CoopNotice.paragraphs[19].runs[2].text = self.loan_number
            CoopNotice.paragraphs[19].runs[2].font.name = 'Times New Roman'
            CoopNotice.paragraphs[21].runs[2].text = ''
            CoopNotice.paragraphs[21].runs[2].font.name = 'Times New Roman'
            CoopNotice.save('News/{} - {} - Coop30DayNotice.docx'.format(self.borrower, self.loan_number))
        elif self.loan_type == 'Coop 1st + HELOC':
            Coop1stHELOCNotice.paragraphs[9].runs[2].text = self.contact
            Coop1stHELOCNotice.paragraphs[9].runs[2].font.name = 'Times New Roman'
            Coop1stHELOCNotice.paragraphs[11].runs[6].text = self.contact_email
            Coop1stHELOCNotice.paragraphs[11].runs[6].font.name = 'Times New Roman'
            Coop1stHELOCNotice.paragraphs[15].runs[2].text = todaysdate.strftime("%m/%d/%Y")
            Coop1stHELOCNotice.paragraphs[15].runs[2].font.name = 'Times New Roman'
            Coop1stHELOCNotice.paragraphs[17].runs[2].text = self.borrower
            Coop1stHELOCNotice.paragraphs[17].runs[2].font.name = 'Times New Roman'
            Coop1stHELOCNotice.paragraphs[19].runs[2].text = self.loan_number
            Coop1stHELOCNotice.paragraphs[19].runs[2].font.name = 'Times New Roman'
            Coop1stHELOCNotice.paragraphs[21].runs[2].text = self.heloc_number
            Coop1stHELOCNotice.paragraphs[21].runs[2].font.name = 'Times New Roman'
            Coop1stHELOCNotice.save('News/{} - {} - Coop30DayNotice.docx'.format(self.borrower, self.loan_number))
        elif self.loan_type == 'Coop HELOC':
            CoopHELOCNotice.paragraphs[9].runs[2].text = self.contact
            CoopHELOCNotice.paragraphs[9].runs[2].font.name = 'Times New Roman'
            CoopHELOCNotice.paragraphs[11].runs[6].text = self.contact_email
            CoopHELOCNotice.paragraphs[11].runs[6].font.name = 'Times New Roman'
            CoopHELOCNotice.paragraphs[15].runs[2].text = todaysdate.strftime("%m/%d/%Y")
            CoopHELOCNotice.paragraphs[15].runs[2].font.name = 'Times New Roman'
            CoopHELOCNotice.paragraphs[17].runs[2].text = self.borrower
            CoopHELOCNotice.paragraphs[17].runs[2].font.name = 'Times New Roman'
            CoopHELOCNotice.paragraphs[19].runs[2].text = ''
            CoopHELOCNotice.paragraphs[19].runs[2].font.name = 'Times New Roman'
            CoopHELOCNotice.paragraphs[21].runs[2].text = self.heloc_number
            CoopHELOCNotice.paragraphs[21].runs[2].font.name = 'Times New Roman'
            CoopHELOCNotice.save('News/{} - {} - Coop30DayNotice.docx'.format(self.borrower, self.heloc_number))

    def send_notice(self, cc):
        if self.loan_type == 'CEMA':
            message = outlook.CreateItem(0)
            message.to = self.contact_email
            message.CC = cc
            message.Subject = '{} - {} - Initial Notice'.format(self.borrower, self.loan_number)
            message.HTMLBody = CemaEmail
            message.Attachments.Add('{}/News/{} - {} - Invoice.pdf'.format(CurrentDirectory, self.borrower, self.loan_number))
            message.Attachments.Add('{}/News/{} - {} - CEMA30DayNotice.pdf'.format(CurrentDirectory, self.borrower, self.loan_number))
            message.Send()
            window['Output'].update('CEMA Initial Notice sent to:')
            window['Output2'].update('{} - {}'.format(self.contact_email, datetime.datetime.now().strftime("%m/%d/%Y at %I:%M:%S %p")))
        elif self.loan_type == 'CEMA HELOC':
            message = outlook.CreateItem(0)
            message.to = self.contact_email
            message.CC = cc
            message.Subject = '{} - {} - Initial Notice'.format(self.borrower, self.heloc_number)
            message.HTMLBody = CemaEmail
            message.Attachments.Add('{}/News/{} - {} - Invoice.pdf'.format(CurrentDirectory, self.borrower, self.heloc_number))
            message.Attachments.Add('{}/News/{} - {} - CEMA30DayNotice.pdf'.format(CurrentDirectory, self.borrower, self.heloc_number))
            message.Send()
            window['Output'].update('CEMA Initial Notice sent to:')
            window['Output2'].update('{} - {}'.format(self.contact_email, datetime.datetime.now().strftime("%m/%d/%Y at %I:%M:%S %p")))
        elif self.loan_type == 'Coop':
            message = outlook.CreateItem(0)
            message.to = self.contact_email
            message.CC = cc
            message.Subject = '{} - {} - Initial Notice'.format(self.borrower, self.loan_number)
            message.HTMLBody = CoopEmail
            message.Attachments.Add('{}/News/{} - {} - Invoice.pdf'.format(CurrentDirectory, self.borrower, self.loan_number))
            message.Attachments.Add('{}/News/{} - {} - Coop30DayNotice.pdf'.format(CurrentDirectory, self.borrower, self.loan_number))
            message.Send()
            window['Output'].update('Coop Initial Notice sent to:')
            window['Output2'].update('{} - {}'.format(self.contact_email, datetime.datetime.now().strftime("%m/%d/%Y at %I:%M:%S %p")))
        elif self.loan_type == 'Coop 1st + HELOC':
            message = outlook.CreateItem(0)
            message.to = self.contact_email
            message.CC = cc
            message.Subject = '{} - {} - Initial Notice'.format(self.borrower, self.loan_number)
            message.HTMLBody = CoopEmail
            message.Attachments.Add('{}/News/{} - {} - Invoice.pdf'.format(CurrentDirectory, self.borrower, self.loan_number))
            message.Attachments.Add('{}/News/{} - {} - Coop30DayNotice.pdf'.format(CurrentDirectory, self.borrower, self.loan_number))
            message.Send()
            window['Output'].update('Coop Initial Notice sent to:')
            window['Output2'].update('{} - {}'.format(self.contact_email, datetime.datetime.now().strftime("%m/%d/%Y at %I:%M:%S %p")))
        elif self.loan_type == 'Coop HELOC':
            message = outlook.CreateItem(0)
            message.to = self.contact_email
            message.CC = cc
            message.Subject = '{} - {} - Initial Notice'.format(self.borrower, self.loan_number)
            message.HTMLBody = CoopEmail
            message.Attachments.Add('{}/News/{} - {} - Invoice.pdf'.format(CurrentDirectory, self.borrower, self.loan_number))
            message.Attachments.Add('{}/News/{} - {} - Coop30DayNotice.pdf'.format(CurrentDirectory, self.borrower, self.loan_number))
            message.Send()
            window['Output'].update('Coop Initial Notice sent to:')
            window['Output2'].update('{} - {}'.format(self.contact_email, datetime.datetime.now().strftime("%m/%d/%Y at %I:%M:%S %p")))



def convert_to_pdf(input_doc, output_pdf):
    tempdoc = word.Documents.Open(os.path.abspath(input_doc))
    tempdoc.SaveAs(os.path.abspath(output_pdf), FileFormat = 17)
    tempdoc.Close()


# Event Loop to process "events" and get the "values" of the inputs
while True:
    event, values = window.read()
    if event == sg.WIN_CLOSED or event == 'Exit':
        word.Quit()
        break
    if event == 'Enter':         
        # Assigning the values taken from the fields to our variable class
        payoff = Payoffs(values['-loannumber-'], 
                         values['-helocnumber-'], 
                         values['-borrower-'], 
                         values['LoanType'], 
                         values['-contact-'], 
                         values['-contactemail-'], 
                         values['-phonenumber-'], 
                         values['-residents-'], 
                         values['-streetaddress-'], 
                         values['-citystatezip-'])

        # Determining Loan Type
        if values['LoanType'] == 'CEMA':
            payoff.add_to_excel()
            payoff.make_invoice(payoff.loan_number, 250)
            payoff.make_initial_notice()
            payoff.make_update_sheet(payoff.loan_number)
            payoff.make_scheduling_form(payoff.loan_number)
            convert_to_pdf('News/{} - {} - Invoice.docx'.format(payoff.borrower, payoff.loan_number), 'News/{} - {} - Invoice.pdf'.format(payoff.borrower, payoff.loan_number))
            convert_to_pdf('News/{} - {} - CEMA30DayNotice.docx'.format(payoff.borrower, payoff.loan_number), 'News/{} - {} - CEMA30DayNotice.pdf'.format(payoff.borrower, payoff.loan_number))
            convert_to_pdf('News/{} - {} - UpdateSheet.docx'.format(payoff.borrower, payoff.loan_number), 'News/{} - {} - UpdateSheet.pdf'.format(payoff.borrower, payoff.loan_number))
            payoff.send_notice('chrisheyer0@gmail.com')
        elif values['LoanType'] == 'CEMA HELOC':
            payoff.add_to_excel()
            payoff.make_invoice(payoff.heloc_number, 250)
            payoff.make_initial_notice()
            payoff.make_update_sheet(payoff.heloc_number)
            payoff.make_scheduling_form(payoff.heloc_number)
            convert_to_pdf('News/{} - {} - Invoice.docx'.format(payoff.borrower, payoff.heloc_number), 'News/{} - {} - Invoice.pdf'.format(payoff.borrower, payoff.heloc_number))
            convert_to_pdf('News/{} - {} - CEMA30DayNotice.docx'.format(payoff.borrower, payoff.heloc_number), 'News/{} - {} - CEMA30DayNotice.pdf'.format(payoff.borrower, payoff.heloc_number))
            convert_to_pdf('News/{} - {} - UpdateSheet.docx'.format(payoff.borrower, payoff.heloc_number), 'News/{} - {} - UpdateSheet.pdf'.format(payoff.borrower, payoff.heloc_number))
            payoff.send_notice('chrisheyer0@gmail.com')
        elif values['LoanType'] == 'Coop':
            payoff.add_to_excel()
            payoff.make_invoice(payoff.loan_number, 250)
            payoff.make_initial_notice()
            payoff.make_update_sheet(payoff.loan_number)
            payoff.make_scheduling_form(payoff.loan_number)
            convert_to_pdf('News/{} - {} - Invoice.docx'.format(payoff.borrower, payoff.loan_number), 'News/{} - {} - Invoice.pdf'.format(payoff.borrower, payoff.loan_number))
            convert_to_pdf('News/{} - {} - Coop30DayNotice.docx'.format(payoff.borrower, payoff.loan_number), 'News/{} - {} - Coop30DayNotice.pdf'.format(payoff.borrower, payoff.loan_number))
            convert_to_pdf('News/{} - {} - UpdateSheet.docx'.format(payoff.borrower, payoff.loan_number), 'News/{} - {} - UpdateSheet.pdf'.format(payoff.borrower, payoff.loan_number))
            payoff.send_notice('chrisheyer0@gmail.com')
        elif values['LoanType'] == 'Coop 1st + HELOC':
            payoff.add_to_excel()
            payoff.make_invoice('{} & {}'.format(payoff.loan_number, payoff.heloc_number), 400)
            payoff.make_initial_notice()
            payoff.make_update_sheet(payoff.loan_number)
            payoff.make_scheduling_form('{} & {}'.format(payoff.loan_number, payoff.heloc_number))
            convert_to_pdf('News/{} - {} - Invoice.docx'.format(payoff.borrower, payoff.loan_number), 'News/{} - {} - Invoice.pdf'.format(payoff.borrower, payoff.loan_number))
            convert_to_pdf('News/{} - {} - Coop30DayNotice.docx'.format(payoff.borrower, payoff.loan_number), 'News/{} - {} - Coop30DayNotice.pdf'.format(payoff.borrower, payoff.loan_number))
            convert_to_pdf('News/{} - {} - UpdateSheet.docx'.format(payoff.borrower, payoff.loan_number), 'News/{} - {} - UpdateSheet.pdf'.format(payoff.borrower, payoff.loan_number))
            payoff.send_notice('chrisheyer0@gmail.com')
        elif values['LoanType'] == 'Coop HELOC':
            payoff.add_to_excel()
            payoff.make_invoice(payoff.heloc_number, 375)
            payoff.make_initial_notice()
            payoff.make_update_sheet(payoff.heloc_number)
            payoff.make_scheduling_form(payoff.heloc_number)
            convert_to_pdf('News/{} - {} - Invoice.docx'.format(payoff.borrower, payoff.heloc_number), 'News/{} - {} - Invoice.pdf'.format(payoff.borrower, payoff.heloc_number))
            convert_to_pdf('News/{} - {} - Coop30DayNotice.docx'.format(payoff.borrower, payoff.heloc_number), 'News/{} - {} - Coop30DayNotice.pdf'.format(payoff.borrower, payoff.heloc_number))
            convert_to_pdf('News/{} - {} - UpdateSheet.docx'.format(payoff.borrower, payoff.heloc_number), 'News/{} - {} - UpdateSheet.pdf'.format(payoff.borrower, payoff.heloc_number))
            payoff.send_notice('chrisheyer0@gmail.com')

    if event == 'Clear':
        window['-borrower-'].update('')
        window['-loannumber-'].update('')
        window['-helocnumber-'].update('')
        window['-contact-'].update('')
        window['-contactemail-'].update('')
        window['-phonenumber-'].update('')
        window['-residents-'].update('')
        window['-streetaddress-'].update('')
        window['-citystatezip-'].update('')
        window['Output2'].update('')
        window['Output'].update('Fields Cleared')

window.close()

